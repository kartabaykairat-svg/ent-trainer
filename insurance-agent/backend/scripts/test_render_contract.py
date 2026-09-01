from pathlib import Path
from docxtpl import DocxTemplate

BASE = Path(__file__).resolve().parent.parent

ctx = {
    "contract_number": "12345",
    "contract_city": "Алматы",
    "contract_year": "2026",
    "contract_year_short": "26",
    "contract_date_kz": "15 қыркүйек",
    "contract_date_day": "15",
    "contract_date_month_ru": "сентября",
    "contract_date_year": "2026",
    "second_insurer": True,
    "c1": {
        "full_name": "Иванов Иван Иванович",
        "birth_date": "01.01.1965",
        "iin": "650101300123",
        "residential_address": "Казахстан, г. Алматы, ул. Абая, 10, кв. 5",
        "doc_type_label_kz": "жеке куәлігі",
        "doc_type_label_ru": "удостоверение личности",
        "doc_series": "22",
        "doc_number": "1234567",
        "doc_issue_day": "12",
        "doc_issue_month_kz": " қаңтар",
        "doc_issue_month_ru": " января",
        "doc_issue_year": "2020",
        "doc_issued_by": "МВД РК",
        "phone": "+7 701 123 45 67",
        "email": "ivanov@example.com",
        "bank_name": "АО Halyk Bank",
        "bank_account": "KZ123456789012345678",
    },
    "c2": {
        "full_name": "Иванова Мария Петровна",
        "birth_date": "02.02.1968",
        "iin": "680202400456",
        "residential_address": "Казахстан, г. Алматы, ул. Абая, 10, кв. 5",
        "doc_type_label_kz": "жеке куәлігі",
        "doc_type_label_ru": "удостоверение личности",
        "doc_series": "22",
        "doc_number": "7654321",
        "doc_issue_day": "12",
        "doc_issue_month_kz": " қаңтар",
        "doc_issue_month_ru": " января",
        "doc_issue_year": "2020",
        "doc_issued_by": "МВД РК",
        "phone": "+7 701 765 43 21",
        "email": "ivanova@example.com",
        "bank_name": "АО Halyk Bank",
        "bank_account": "KZ876543210987654321",
    },
    "premium": {
        "total": "10 000 000", "total_words": "десять миллионов",
        "other_org_c1": "3 000 000", "other_org_c1_words": "три миллиона",
        "other_org_c2": "2 000 000", "other_org_c2_words": "два миллиона",
        "enpf_c1": "4 000 000", "enpf_c1_words": "четыре миллиона",
        "enpf_c2": "1 500 000", "enpf_c2_words": "один миллион пятьсот тысяч",
        "own_c1": "500 000", "own_c1_words": "пятьсот тысяч",
        "own_c2": "0", "own_c2_words": "ноль",
        "premium_c1": "6 500 000",
        "premium_c2": "3 500 000",
    },
    "first_payment": {"c1": "50 000", "c1_words": "пятьдесят тысяч", "c2": "40 000", "c2_words": "сорок тысяч"},
    "guarantee": {
        "years": "15",
        "c1_from": "2026", "c1_to": "2041",
        "c2_from": "2028", "c2_to": "2043",
        "c1_from_day": "15", "c1_from_month": "сентября", "c1_from_year": "2026",
        "c1_to_day": "15", "c1_to_month": "сентября", "c1_to_year": "2041",
        "c2_from_day": "20", "c2_from_month": "марта", "c2_from_year": "2028",
        "c2_to_day": "20", "c2_to_month": "марта", "c2_to_year": "2043",
    },
    "death_benefit": "200 000", "death_benefit_words": "двести тысяч",
    "beneficiary": {"full_name": "", "address": "", "iin": "", "document": ""},
    "schedule": [
        {"date_c1": "15.10.2026", "amount_c1": "50 000", "buyout_c1": "9 500 000",
         "date_c2": "20.04.2028", "amount_c2": "40 000", "buyout_c2": "3 400 000"},
        {"date_c1": "15.11.2026", "amount_c1": "50 000", "buyout_c1": "9 450 000",
         "date_c2": "20.05.2028", "amount_c2": "40 000", "buyout_c2": "3 350 000"},
    ],
    "summary": {
        "age_c1": "61", "age_c1_label": "лет", "gender_c1": "мужской",
        "age_c2": "58", "age_c2_label": "лет", "gender_c2": "женский",
        "contract_basis": "ст. 225 Социального кодекса РК",
        "term_c1": "пожизненно", "term_c2": "пожизненно",
        "effective_rate": "4.5%", "indexation_rate": "7%",
        "cost_ratio_premium": "1.5%", "cost_ratio_payment": "3%",
        "pv_factor": "12.34", "pv_factor_costs": "12.10",
    },
}

tpl = DocxTemplate(str(BASE / "templates" / "contract_template.docx"))
tpl.render(ctx)
out = BASE / "storage" / "generated" / "test_contract.docx"
out.parent.mkdir(parents=True, exist_ok=True)
tpl.save(str(out))
print("rendered ->", out)

# Validate: scan for leftover placeholders
import docx as docxlib
import re
dd = docxlib.Document(str(out))
issues = []
patterns = [r"\{\{.*?\}\}", r"\{%.*?%\}", r"undefined", r"\bNone\b", r"\[ФИО\]"]
def scan_text(t, loc):
    for pat in patterns:
        if re.search(pat, t):
            issues.append((loc, pat, t[:120]))

for i, p in enumerate(dd.paragraphs):
    scan_text(p.text, f"para{i}")
for ti, t in enumerate(dd.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            scan_text(cell.text, f"table{ti}r{ri}c{ci}")

if issues:
    print(f"FOUND {len(issues)} ISSUES:")
    for loc, pat, txt in issues:
        print(" -", loc, pat, repr(txt))
else:
    print("No leftover placeholders found. OK")
