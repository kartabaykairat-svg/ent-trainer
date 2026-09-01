"""OCR: turns an uploaded file (image, PDF or DOCX) into plain text.

Images and scanned PDFs go through Tesseract (rus+kaz+eng language packs).
Native DOCX/PDF text layers are read directly when present (much more
reliable than OCR when available). This module never invents text - if
nothing can be extracted it returns an empty string and the caller must
surface that to the manager rather than guessing.
"""
import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from PIL import Image

from app.config import get_settings

logger = logging.getLogger(__name__)

IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}

TESSERACT_LANGS = "rus+kaz+eng"


def _configure_tesseract():
    settings = get_settings()
    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd


def ocr_image_bytes(data: bytes) -> tuple[str, float]:
    """Returns (text, mean_confidence 0-100)."""
    _configure_tesseract()
    image = Image.open(io.BytesIO(data))
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    try:
        tsv = pytesseract.image_to_data(image, lang=TESSERACT_LANGS, output_type=pytesseract.Output.DICT)
    except pytesseract.TesseractError as exc:
        logger.warning("tesseract failed: %s", exc)
        return "", 0.0

    words = []
    confidences = []
    for text, conf in zip(tsv["text"], tsv["conf"]):
        text = text.strip()
        conf = float(conf) if str(conf).strip() not in ("", "-1") else None
        if text:
            words.append(text)
        if conf is not None and conf >= 0:
            confidences.append(conf)
    full_text = pytesseract.image_to_string(image, lang=TESSERACT_LANGS)
    mean_conf = sum(confidences) / len(confidences) if confidences else 0.0
    return full_text, mean_conf


def ocr_pdf_bytes(data: bytes) -> tuple[str, float]:
    """Uses the PDF's own text layer when present; otherwise rasterizes
    each page and runs Tesseract on it."""
    doc = fitz.open(stream=data, filetype="pdf")
    text_layers = []
    has_text = False
    for page in doc:
        t = page.get_text().strip()
        if t:
            has_text = True
        text_layers.append(t)

    if has_text:
        return "\n".join(text_layers), 95.0  # native text layer, treat as high confidence

    # Scanned PDF: rasterize + OCR each page.
    all_text = []
    all_conf = []
    for page in doc:
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        text, conf = ocr_image_bytes(img_bytes)
        all_text.append(text)
        all_conf.append(conf)
    mean_conf = sum(all_conf) / len(all_conf) if all_conf else 0.0
    return "\n".join(all_text), mean_conf


def read_docx_text(data: bytes) -> tuple[str, float]:
    import docx

    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts), 98.0


def extract_text(filename: str, data: bytes, content_type: str = "") -> tuple[str, float]:
    """Dispatches on file extension. Returns (text, confidence 0-100)."""
    suffix = Path(filename).suffix.lower()
    try:
        if suffix in IMAGE_SUFFIXES:
            return ocr_image_bytes(data)
        if suffix in PDF_SUFFIXES:
            return ocr_pdf_bytes(data)
        if suffix in DOCX_SUFFIXES:
            return read_docx_text(data)
    except Exception:
        logger.exception("OCR/text extraction failed for %s", filename)
        return "", 0.0
    logger.warning("Unsupported file type for OCR: %s", filename)
    return "", 0.0
