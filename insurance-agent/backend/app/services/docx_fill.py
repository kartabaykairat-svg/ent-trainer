"""Turns a validated ClientData object into the two filled DOCX files
(contract + power of attorney), using the annotated templates produced by
scripts/annotate_contract.py / scripts/annotate_poa.py and documented in
templates/field_mapping.json.

Two safety nets are enforced before a document ever reaches the manager:
  * pre_generation_validate() checks the *data* (section 13 checklist)
    before rendering is even attempted.
  * scan_for_leftover_placeholders() checks the *rendered file* afterwards
    for anything that would mean a tag failed to resolve (jinja braces,
    'undefined', 'None', 'null', '{{name}}', '[ФИО]', ...). Either check
    failing blocks generation.
"""
import re
from dataclasses import dataclass
from pathlib import Path

import docx
from docxtpl import DocxTemplate

from app.config import get_settings
from app.schemas import ClientData, PersonData
from app.services.text_utils import (
    month_kz,
    month_ru_genitive,
    number_to_words_ru,
    parse_ddmmyyyy,
)

DOC_TYPE_LABELS = {
    "id_card": ("жеке куәлігі", "удостоверение личности"),
    "passport": ("төлқұжаты", "паспорт"),
    "residence_permit": ("тұруға ықтиярхаты", "вид на жительство"),
    "stateless_id": ("азаматтығы жоқ адамның куәлігі", "удостоверение лица без гражданства"),
}


def _doc_labels(doc_type: str) -> tuple[str, str]:
    return DOC_TYPE_LABELS.get(doc_type, ("құжаты", "документ"))


def _split_date(value: str) -> tuple[str, str, str, str, str]:
    """DD.MM.YYYY -> (day, month_kz, month_ru_with_leading_space, month_ru_bare, year)"""
    parsed = parse_ddmmyyyy(value)
    if not parsed:
        return "", "", "", "", ""
    d, m, y = parsed
    return f"{d:02d}", month_kz(m), f" {month_ru_genitive(m)}", month_ru_genitive(m), f"{y:04d}"


def _person_context(person: PersonData) -> dict:
    day, mkz, mru_lead, mru, year = _split_date(person.document.issue_date)
    label_kz, label_ru = _doc_labels(person.document.type)
    return {
        "full_name": person.full_name,
        "birth_date": person.birth_date,
        "iin": person.iin,
        "residential_address": person.residential_address or person.registration_address,
        "doc_type_label_kz": label_kz,
        "doc_type_label_ru": label_ru,
        "doc_series": person.document.series,
        "doc_number": person.document.number,
        "doc_issue_day": day,
        "doc_issue_month_kz": mkz,
        "doc_issue_month_ru": mru_lead,
        "doc_issue_year": year,
        "doc_issued_by": person.document.issued_by,
        "phone": person.phone,
        "email": person.email,
        "bank_name": person.bank.name,
        "bank_account": person.bank.iban or person.bank.account,
    }


def build_contract_context(client: ClientData) -> dict:
    ins = client.insurance
    day, _mkz, mru_lead, mru, year = _split_date(ins.contract_date)

    def money(v: str) -> dict:
        return {"amount": v or "0", "words": number_to_words_ru(v) or "ноль"}

    premium = {
        "total": ins.premium_own_c1 or "",  # placeholder overwritten below with real total
        "total_words": "",
        "other_org_c1": ins.premium_other_org_c1, "other_org_c1_words": number_to_words_ru(ins.premium_other_org_c1),
        "other_org_c2": ins.premium_other_org_c2, "other_org_c2_words": number_to_words_ru(ins.premium_other_org_c2),
        "enpf_c1": ins.premium_enpf_c1, "enpf_c1_words": number_to_words_ru(ins.premium_enpf_c1),
        "enpf_c2": ins.premium_enpf_c2, "enpf_c2_words": number_to_words_ru(ins.premium_enpf_c2),
        "own_c1": ins.premium_own_c1, "own_c1_words": number_to_words_ru(ins.premium_own_c1),
        "own_c2": ins.premium_own_c2, "own_c2_words": number_to_words_ru(ins.premium_own_c2),
    }
    total_c1 = _sum_amounts(ins.premium_other_org_c1, ins.premium_enpf_c1, ins.premium_own_c1)
    total_c2 = _sum_amounts(ins.premium_other_org_c2, ins.premium_enpf_c2, ins.premium_own_c2)
    total = _sum_amounts(total_c1, total_c2)
    premium["total"] = _fmt_amount(total)
    premium["total_words"] = number_to_words_ru(total)
    premium["premium_c1"] = _fmt_amount(total_c1)
    premium["premium_c2"] = _fmt_amount(total_c2)

    g1f = _split_date(ins.guarantee_c1_from)
    g1t = _split_date(ins.guarantee_c1_to)
    g2f = _split_date(ins.guarantee_c2_from)
    g2t = _split_date(ins.guarantee_c2_to)

    calc = client.calculation
    age_c1 = calc.age_c1 if calc else None
    age_c2 = calc.age_c2 if calc else None

    schedule = [
        {
            "date_c1": s.date_c1, "amount_c1": s.amount_c1, "buyout_c1": s.buyout_c1,
            "date_c2": s.date_c2, "amount_c2": s.amount_c2, "buyout_c2": s.buyout_c2,
        }
        for s in client.schedule
    ]

    return {
        "contract_number": ins.contract_number,
        "contract_city": ins.contract_city,
        "contract_year": year,
        "contract_year_short": year[-2:] if year else "",
        "contract_date_kz": f"{day} {month_kz(parse_ddmmyyyy(ins.contract_date)[1]) if parse_ddmmyyyy(ins.contract_date) else ''}",
        "contract_date_day": day,
        "contract_date_month_ru": mru,
        "contract_date_year": year,
        "second_insurer": client.second_insurer,
        "c1": _person_context(client.c1),
        "c2": _person_context(client.c2) if client.second_insurer else _person_context(PersonData()),
        "premium": premium,
        "first_payment": {
            "c1": ins.first_payment_c1, "c1_words": number_to_words_ru(ins.first_payment_c1),
            "c2": ins.first_payment_c2, "c2_words": number_to_words_ru(ins.first_payment_c2),
        },
        "guarantee": {
            "years": ins.guarantee_years,
            "c1_from": g1f[4], "c1_to": g1t[4], "c2_from": g2f[4], "c2_to": g2t[4],
            "c1_from_day": g1f[0], "c1_from_month": g1f[3], "c1_from_year": g1f[4],
            "c1_to_day": g1t[0], "c1_to_month": g1t[3], "c1_to_year": g1t[4],
            "c2_from_day": g2f[0], "c2_from_month": g2f[3], "c2_from_year": g2f[4],
            "c2_to_day": g2t[0], "c2_to_month": g2t[3], "c2_to_year": g2t[4],
        },
        "death_benefit": ins.death_benefit, "death_benefit_words": number_to_words_ru(ins.death_benefit),
        "beneficiary": client.beneficiary.model_dump(),
        "schedule": schedule,
        "summary": {
            "age_c1": age_c1 or "", "age_c1_label": _age_label(age_c1),
            "age_c2": age_c2 or "", "age_c2_label": _age_label(age_c2),
            "gender_c1": client.c1.gender, "gender_c2": client.c2.gender,
            "contract_basis": "ст. 225 Социального кодекса Республики Казахстан",
            "term_c1": "пожизненно", "term_c2": "пожизненно" if client.second_insurer else "",
            "effective_rate": "", "indexation_rate": f"{ins.indexation_rate} %" if ins.indexation_confirmed else "",
            "cost_ratio_premium": "1,5", "cost_ratio_payment": "3",
            "pv_factor": "", "pv_factor_costs": "",
        },
    }


def _age_label(age: int | None) -> str:
    if age is None:
        return ""
    n10, n100 = age % 10, age % 100
    if 11 <= n100 <= 14:
        return "лет"
    if n10 == 1:
        return "год"
    if 2 <= n10 <= 4:
        return "года"
    return "лет"


def _sum_amounts(*values: str) -> str:
    total = 0
    for v in values:
        digits = re.sub(r"[^\d]", "", v or "")
        if digits:
            total += int(digits)
    return str(total)


def _fmt_amount(v: str) -> str:
    digits = re.sub(r"[^\d]", "", v or "")
    if not digits:
        return "0"
    return f"{int(digits):,}".replace(",", " ")


def build_poa_context(client: ClientData) -> dict:
    ins = client.insurance
    day, _mkz, mru_lead, mru, year = _split_date(ins.contract_date)
    rep = client.representative if client.representative_override_enabled else None
    return {
        "poa_city": ins.contract_city,
        "poa_date_day": day,
        "poa_date_month_ru": mru,
        "poa_date_year": year,
        "c1": {
            "full_name": client.c1.full_name,
            "birth_date": client.c1.birth_date,
            "iin": client.c1.iin,
            "registration_address": client.c1.registration_address or client.c1.residential_address,
        },
        "representative": (rep.model_dump() if rep else {}),
    }


def render_docx(template_path: Path, context: dict, output_path: Path) -> None:
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(output_path))


_PLACEHOLDER_PATTERNS = [
    re.compile(r"\{\{.*?\}\}"),
    re.compile(r"\{%.*?%\}"),
    re.compile(r"\bundefined\b", re.IGNORECASE),
    re.compile(r"\bNone\b"),
    re.compile(r"\bnull\b", re.IGNORECASE),
    re.compile(r"\[ФИО\]"),
    re.compile(r"\[Ф\.?И\.?О\.?\]", re.IGNORECASE),
]


def scan_for_leftover_placeholders(docx_path: Path) -> list[str]:
    d = docx.Document(str(docx_path))
    issues: set[str] = set()

    def scan(text: str):
        for pat in _PLACEHOLDER_PATTERNS:
            m = pat.search(text)
            if m:
                issues.add(f"Незаполненный шаблонный тег в документе: {m.group()!r}")

    for p in d.paragraphs:
        scan(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                scan(cell.text)
    return sorted(issues)
