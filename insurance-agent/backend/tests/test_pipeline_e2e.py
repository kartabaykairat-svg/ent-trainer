"""Full pipeline smoke test: login -> create client -> upload real OCR'd
images -> review merged data -> set insurance params -> confirm -> generate
contract+POA DOCX -> download -> assert no leftover placeholders and that
key client values actually appear in the rendered files.

Run with: pytest tests/test_pipeline_e2e.py -v
"""
import io
from pathlib import Path

import docx
import pytest
from fastapi.testclient import TestClient

FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture()
def client(tmp_path, monkeypatch):
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{tmp_path}/test.db")
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path / "storage"))
    from app.config import get_settings
    get_settings.cache_clear()

    from app.main import app
    with TestClient(app) as c:
        yield c
    get_settings.cache_clear()


def auth_headers(client) -> dict:
    r = client.post("/api/auth/login", json={"username": "manager", "password": "manager123"})
    assert r.status_code == 200, r.text
    return {"Authorization": f"Bearer {r.json()['access_token']}"}


def test_full_pipeline(client):
    headers = auth_headers(client)

    r = client.post("/api/clients", json={"second_insurer": False}, headers=headers)
    assert r.status_code == 200, r.text
    client_id = r.json()["id"]

    files = [
        ("files", ("c1_id.png", (FIXTURES / "c1_id.png").read_bytes(), "image/png")),
        ("files", ("c1_address.png", (FIXTURES / "c1_address.png").read_bytes(), "image/png")),
    ]
    r = client.post(f"/api/clients/{client_id}/documents?belongs_to=c1", files=files, headers=headers)
    assert r.status_code == 200, r.text
    upload_result = r.json()
    for u in upload_result["uploaded"]:
        assert "error" not in u, u

    data = upload_result["client"]["data"]
    assert data["c1"]["iin"] == "650101300123"
    assert data["c1"]["birth_date"] == "01.01.1965"
    assert "ИВАНОВ" in data["c1"]["full_name"]
    assert data["c1"]["residential_address"], "address should have been extracted"
    assert data["c1"]["document"]["number"] == "032614460"

    # insurance parameters (manager-entered, required)
    insurance_payload = {
        "contract_number": "12345",
        "contract_date": "15.09.2026",
        "contract_city": "Алматы",
        "premium_other_org_c1": "3000000",
        "premium_enpf_c1": "4000000",
        "premium_own_c1": "500000",
        "first_payment_c1": "50000",
        "payment_periodicity": "ежемесячно",
        "guarantee_years": "15",
        "guarantee_c1_from": "15.09.2026",
        "guarantee_c1_to": "15.09.2041",
        "death_benefit": "200000",
        "indexation_rate": "7",
        "indexation_confirmed": True,
        "bank_name_c1": "АО Halyk Bank",
        "bank_account_c1": "KZ123456789012345678",
    }
    r = client.put(f"/api/clients/{client_id}/insurance", json=insurance_payload, headers=headers)
    assert r.status_code == 200, r.text

    # confirm should fail: c1 phone/email present but nothing blocking except maybe none - let's just confirm
    r = client.post(f"/api/clients/{client_id}/confirm", headers=headers)
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["ok"] is True, body.get("errors")

    r = client.get(f"/api/clients/{client_id}/checklist", headers=headers)
    assert r.status_code == 200
    assert r.json()["passed"] is True, r.json()

    r = client.post(f"/api/clients/{client_id}/generate", headers=headers)
    assert r.status_code == 200, r.text
    assert r.json()["pdf_available"] is True, "LibreOffice should be available in this environment"

    r = client.get(f"/api/clients/{client_id}/download/contract/docx", headers=headers)
    assert r.status_code == 200
    contract_bytes = r.content
    d = docx.Document(io.BytesIO(contract_bytes))
    full_text = "\n".join(p.text for p in d.paragraphs) + "\n".join(
        cell.text for t in d.tables for row in t.rows for cell in row.cells
    )
    assert "ИВАНОВ ИВАН ИВАНОВИЧ" in full_text
    assert "650101300123" in full_text
    assert "{{" not in full_text and "{%" not in full_text
    assert "undefined" not in full_text.lower()

    r = client.get(f"/api/clients/{client_id}/download/contract/pdf", headers=headers)
    assert r.status_code == 200
    assert r.headers["content-type"] == "application/pdf"
    assert r.content[:4] == b"%PDF"

    r = client.get(f"/api/clients/{client_id}/download/poa/docx", headers=headers)
    assert r.status_code == 200
    poa_doc = docx.Document(io.BytesIO(r.content))
    poa_text = "\n".join(p.text for p in poa_doc.paragraphs)
    assert "ИВАНОВ ИВАН ИВАНОВИЧ" in poa_text
    assert "УНДИЗОВА ФАГИЛЯМ" in poa_text  # default representative preserved
    assert "{{" not in poa_text

    r = client.get("/api/dashboard", headers=headers)
    assert r.status_code == 200
    dash = r.json()
    assert dash["contracts_generated"] >= 1
    assert dash["poa_generated"] >= 1

    r = client.get("/api/clients", headers=headers)
    listing = r.json()
    assert listing[0]["masked_iin"] == "650101******"
    assert "1234" not in listing[0]["masked_iin"]
