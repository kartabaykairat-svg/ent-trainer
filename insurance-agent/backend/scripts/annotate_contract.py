"""
One-time template-preparation script.

Reads the ORIGINAL, unmodified insurer template
(templates_source/contract_original.docx) and produces an annotated copy
(templates/contract_template.docx) in which every blank ("____") that is
meant to be filled with client/contract data is replaced with a Jinja2
merge-tag ({{ field }}) understood by docxtpl at render time.

Design rules followed (see field_mapping.json for the human-readable map):
  * No legal wording, article numbers, clause numbering, headings, table
    structure, fonts or page setup are touched.
  * Only runs that already consisted of blank underscores (or the empty
    paragraphs reserved for hand-written values) are edited.
  * Where a blank lived inside a single run, only that run's text is
    replaced (`run.text = ...`) - every sibling run keeps its original
    formatting untouched.
  * Where a blank was split across a few runs (Word track-changes / manual
    edits do this often), the runs are merged into the first run of the
    blank *only* - the label runs around it are never touched.
  * For long, sentence-style clauses that contain many blanks inline
    (the premium clause, the death-benefit clause, etc.) the paragraph is
    rebuilt as a single run from its own text, with the blanks replaced by
    tags - the wording itself is copied verbatim from the source, nothing
    is summarized, shortened or reworded.

Run this script again any time the two source templates change; it always
starts from templates_source/*.docx, never from a previously annotated file.
"""
import re
from pathlib import Path

import docx

BASE = Path(__file__).resolve().parent.parent
SRC = BASE / "templates_source" / "contract_original.docx"
DST = BASE / "templates" / "contract_template.docx"


def merge_runs(paragraph, start=0, end=None):
    """Merge paragraph.runs[start:end+1] into the first run of the range."""
    runs = paragraph.runs
    if not runs:
        return None
    if end is None:
        end = len(runs) - 1
    for i in range(start + 1, end + 1):
        runs[i].text = ""
    return runs[start]


def set_para(paragraph, text):
    """Replace the whole paragraph's visible text with `text` in a single
    run, reusing the first run's formatting. Text is taken from the
    template's own wording with blanks swapped for tags - nothing new is
    invented."""
    r = merge_runs(paragraph)
    if r is None:
        paragraph.add_run(text)
    else:
        r.text = text


def set_run_text(paragraph, idx, text):
    paragraph.runs[idx].text = text


def cell_para(table, row, col, para_idx=0):
    return table.rows[row].cells[col].paragraphs[para_idx]


def distinct_cells(row):
    """Return each physically distinct <w:tc> in a row once, skipping the
    duplicate accessor python-docx returns for the grid positions covered
    by a horizontally-merged (gridSpan) cell."""
    seen = set()
    out = []
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            out.append(c)
    return out


def main():
    d = docx.Document(SRC)
    t0, t1, t2, t3 = d.tables

    # ---- Contract number on the cover page ("сериясы / серия 22 № ___") --
    set_run_text(d.paragraphs[3], 5, "{{ contract_number }}")

    # ---- Appendix 1 header (payment schedule) ----------------------------
    set_para(d.paragraphs[41], "20{{ contract_year_short }}ж. {{ contract_date_kz }}№{{ contract_number }}")
    set_run_text(d.paragraphs[46], 1, "№ {{ contract_number }}")
    set_para(d.paragraphs[47], "от «{{ contract_date_day }}» {{ contract_date_month_ru }} {{ contract_date_year }} г.")

    # ---- Appendix 2 header (заключение) -----------------------------------
    set_para(d.paragraphs[68], "20{{ contract_year_short }}ж. {{ contract_date_kz }}№{{ contract_number }}")
    set_run_text(d.paragraphs[73], 1, "№ {{ contract_number }}")
    set_para(d.paragraphs[74], "от «{{ contract_date_day }}» {{ contract_date_month_ru }} {{ contract_date_year }} г.")

    # ---- "Заключение к договору" number/date line -------------------------
    p79 = d.paragraphs[79]
    set_run_text(p79, 3, "{{ contract_number }}")
    set_run_text(p79, 5, "«{{ contract_date_day }}»______________20{{ contract_year_short }}")

    # ======================================================================
    # TABLE 0 - main body of the contract
    # ======================================================================

    # -- Row 0: place & year of conclusion (KZ col0 / RU col2) -------------
    p = cell_para(t0, 0, 0)
    set_run_text(p, 1, "{{ contract_city }}")
    set_run_text(p, 4, "{{ contract_year }}")
    p = cell_para(t0, 0, 2)
    set_run_text(p, 1, "{{ contract_city }}")
    set_run_text(p, 4, "{{ contract_year }}")

    # -- Row 1: Strakhователь(и) preamble, KZ column (col 0) ----------------
    c1_details_kz = (
        "{{ c1.birth_date }}, жеке сәйкестендіру нөмірі {{ c1.iin }}, "
        "{{ c1.residential_address }}, {{ c1.doc_type_label_kz }} сериясы {{ c1.doc_series }} "
        "№ {{ c1.doc_number }}, «{{ c1.doc_issue_day }}»{{ c1.doc_issue_month_kz }} "
        "{{ c1.doc_issue_year }} жылы берілген),"
    )
    c2_details_kz = (
        "{% if second_insurer %}{{ c2.birth_date }}, жеке сәйкестендіру нөмірі {{ c2.iin }}, "
        "{{ c2.residential_address }}, {{ c2.doc_type_label_kz }} сериясы {{ c2.doc_series }} "
        "№ {{ c2.doc_number }}, «{{ c2.doc_issue_day }}»{{ c2.doc_issue_month_kz }} "
        "{{ c2.doc_issue_year }} жылы берілген),{% endif %}"
    )
    cellkz = t0.rows[1].cells[0]
    p2 = cellkz.paragraphs[2]
    merge_runs(p2, 2, 3)
    set_run_text(p2, 2, " {{ c1.full_name }}")
    set_para(cellkz.paragraphs[4], c1_details_kz)
    set_para(cellkz.paragraphs[7], "{% if second_insurer %}{{ c2.full_name }}{% endif %}")
    set_para(cellkz.paragraphs[9], c2_details_kz)

    # -- Row 1: Strakhователь(и) preamble, RU column (col 2) ---------------
    c1_details_ru = (
        "{{ c1.birth_date }}, индивидуальный идентификационный номер {{ c1.iin }}, "
        "место жительства {{ c1.residential_address }}, {{ c1.doc_type_label_ru }}, серии {{ c1.doc_series }} "
        "№ {{ c1.doc_number }}, выданный {{ c1.doc_issued_by }} «{{ c1.doc_issue_day }}»"
        "{{ c1.doc_issue_month_ru }} {{ c1.doc_issue_year }} года),"
    )
    c2_details_ru = (
        "{% if second_insurer %}{{ c2.birth_date }}, индивидуальный идентификационный номер {{ c2.iin }}, "
        "место жительства {{ c2.residential_address }}, {{ c2.doc_type_label_ru }}, серии {{ c2.doc_series }} "
        "№ {{ c2.doc_number }}, выданный {{ c2.doc_issued_by }} «{{ c2.doc_issue_day }}»"
        "{{ c2.doc_issue_month_ru }} {{ c2.doc_issue_year }} года),{% endif %}"
    )
    cellru = t0.rows[1].cells[2]
    set_run_text(cellru.paragraphs[2], 0, "{{ c1.full_name }} ")
    set_para(cellru.paragraphs[4], c1_details_ru)
    set_para(cellru.paragraphs[7], "{% if second_insurer %}{{ c2.full_name }} {% endif %}")
    set_para(cellru.paragraphs[9], c2_details_ru)

    # ======================================================================
    # Row 7 - premium / payment clauses (long numbered sentences, both
    # columns). Rebuilt verbatim from the source wording, blanks -> tags.
    # ======================================================================
    kz7 = t0.rows[7].cells[0]
    set_para(kz7.paragraphs[0],
        "2. Сақтанушының (Сақтанушылардың)  сақтандыру сыйлықақысының мөлшері "
        "{{ premium.total_words }} ({{ premium.total }}) теңгені құрайды және басқа сақтандыру "
        "ұйымынан {{ premium.other_org_c1_words }} ({{ premium.other_org_c1 }}) теңге және "
        "{% if second_insurer %}{{ premium.other_org_c2_words }} ({{ premium.other_org_c2 }}) теңге{% else %}0 теңге{% endif %} "
        "(екінші сақтанушы болған кезде) мөлшерінде сатып алу сомасынан және бірыңғай жинақтаушы "
        "зейнетақы қорынан {{ premium.enpf_c1_words }} ({{ premium.enpf_c1 }}) теңге және "
        "{% if second_insurer %}{{ premium.enpf_c2_words }} ({{ premium.enpf_c2 }}) теңге{% else %}0 теңге{% endif %} "
        "(екінші сақтанушы болған кезде) мөлшерінде зейнетақы жинақтарынан, және Сақтанушының "
        "(-лардың) {{ premium.own_c1_words }} ({{ premium.own_c1 }}) теңге, және "
        "{% if second_insurer %}{{ premium.own_c2_words }} ({{ premium.own_c2 }}) теңге{% else %}0 теңге{% endif %} "
        "(екінші сақтанушы болған кезде) мөлшерінде меншікті қаражатынан тұрады.")
    set_para(kz7.paragraphs[4],
        "Сақтандырылушыға бірінші ай сайынғы сақтандыру төлемінің мөлшері  "
        "{{ first_payment.c1_words }} ({{ first_payment.c1 }}) теңгені және "
        "{% if second_insurer %}{{ first_payment.c2_words }} ({{ first_payment.c2 }}) теңгені{% else %}0 теңгені{% endif %} "
        "(екінші сақтандырылушы болған кезде) құрайды. ")
    set_para(kz7.paragraphs[6],
        "7.  Кепілдік берілген сақтандыру төлемдерін (бар болса ) жүзеге асыру кезеңі "
        "{{ guarantee.c1_from }} {{ guarantee.c1_to }} "
        "{% if second_insurer %}және {{ guarantee.c2_from }} - {{ guarantee.c2_to }}{% endif %} "
        "аралығында {{ guarantee.years }} жыл (екінші сақтанушы бар болса) құрайды.")
    set_para(kz7.paragraphs[10],
        "10. Сақтанушы (Сақтанушылар) және (немесе) сақтандырылушылар қайтыс болған жағдайда, "
        "Сақтандырушы отбасына не жерлеуді жүзеге асырған адамға {{ death_benefit_words }} "
        "({{ death_benefit }}) теңге мөлшерінде, бірақ тиісті қаржы жылына арналған республикалық "
        "бюджет туралы заңда белгіленген айлық есептік көрсеткіштің 35 еселенген мөлшерінен кем "
        "емес жерлеуге бір жолғы төлем түрінде әрбір Сақтанушыға (екінші сақтанушы болған кезде) "
        "сақтандыру төлемін (төлемдерін) жүзеге асырады.")
    set_para(kz7.paragraphs[13],
        "{% if beneficiary.full_name %}{{ beneficiary.full_name }} (тегі, аты, жөні (бар болса)), "
        "{{ beneficiary.address }}{% else %}____________________________________________ "
        "(тегі, аты, жөні (бар болса)), ____________________________________________{% endif %}")
    set_para(kz7.paragraphs[15],
        "{% if beneficiary.full_name %}{{ beneficiary.iin }} (жеке сәйкестендіру нөмірі), "
        "{{ beneficiary.document }} (жеке басын куәландыратын құжат) (бірнеше алушылар болған кезде "
        "деректер әрбір алушы бойынша жеке көрсетіледі) болып табылады.{% else %}"
        "____________________________________________ (жеке сәйкестендіру нөмірі), "
        "____________________________________________ (жеке басын куәландыратын құжат) (бірнеше "
        "алушылар болған кезде деректер әрбір алушы бойынша жеке көрсетіледі) болып табылады.{% endif %}")

    ru7 = t0.rows[7].cells[2]
    set_para(ru7.paragraphs[0],
        "2. Размер страховой премии Страхователя (Страхователей) составляет "
        "{{ premium.total_words }} ({{ premium.total }}) тенге и состоит из выкупной суммы из "
        "другой страховой организации в размере {{ premium.other_org_c1_words }} "
        "({{ premium.other_org_c1 }}) тенге и "
        "{% if second_insurer %}{{ premium.other_org_c2_words }} ({{ premium.other_org_c2 }}) тенге{% else %}0 тенге{% endif %} "
        "(при наличии второго страхователя), пенсионных накоплений из единого накопительного "
        "пенсионного фонда в размере {{ premium.enpf_c1_words }} ({{ premium.enpf_c1 }}) тенге и "
        "{% if second_insurer %}{{ premium.enpf_c2_words }} ({{ premium.enpf_c2 }}) тенге{% else %}0 тенге{% endif %} "
        "(при наличии второго страхователя), и собственных средств Страхователя (-ей) в размере "
        "{{ premium.own_c1_words }} ({{ premium.own_c1 }}) тенге и "
        "{% if second_insurer %}{{ premium.own_c2_words }} ({{ premium.own_c2 }}) тенге{% else %}0 тенге{% endif %} "
        "(при наличии второго страхователя).")
    set_para(ru7.paragraphs[4],
        "Размер первой ежемесячной страховой выплаты застрахованному составляет "
        "{{ first_payment.c1_words }} ({{ first_payment.c1 }}) тенге и "
        "{% if second_insurer %}{{ first_payment.c2_words }} ({{ first_payment.c2 }}) тенге{% else %}0 тенге{% endif %} "
        "(при наличии второго застрахованного). ")
    set_para(ru7.paragraphs[6],
        "7. Период осуществления гарантированных страховых выплат (при наличии) составляет "
        "{{ guarantee.years }} года (лет), с {{ guarantee.c1_from }} года по {{ guarantee.c1_to }} "
        "года{% if second_insurer %} и с {{ guarantee.c2_from }} года по {{ guarantee.c2_to }} года "
        "(при наличии второго страхователя){% endif %}.")
    set_para(ru7.paragraphs[10],
        "10. В случае смерти Страхователя (Страхователей) и (или) застрахованных Страховщик "
        "осуществляет страховую (страховые) выплату (выплаты) в виде единовременной выплаты на "
        "погребение семье либо лицу, осуществившему погребение, в размере {{ death_benefit_words }} "
        "({{ death_benefit }}) тенге, но не менее 35-кратного размера месячного расчетного "
        "показателя, установленного на соответствующий финансовый год законом о республиканском "
        "бюджете каждому Страхователю (при наличии второго страхователя).")
    set_para(ru7.paragraphs[13],
        "{% if beneficiary.full_name %}{{ beneficiary.full_name }} (фамилия, имя, отчество (при его "
        "наличии)), {{ beneficiary.address }}{% else %}____________________________________________ "
        "(фамилия, имя, отчество (при его наличии)), "
        "____________________________________________{% endif %}")
    set_para(ru7.paragraphs[15],
        "{% if beneficiary.full_name %}(адрес места жительства), {{ beneficiary.iin }} "
        "(индивидуальный идентификационный номер), {{ beneficiary.document }} (документ, "
        "удостоверяющий личность) (при наличии нескольких получателей данные указываются по каждому "
        "отдельно.{% else %}(адрес места жительства), ____________________________________________ "
        "(индивидуальный идентификационный номер), ____________________________________________ "
        "(документ, удостоверяющий личность) (при наличии нескольких получателей данные указываются "
        "по каждому отдельно.{% endif %}")

    # ======================================================================
    # TABLE 1 - signatures block (KZ col0 / RU col1), row 4
    # ======================================================================
    kzs = t1.rows[4].cells[0]
    set_para(kzs.paragraphs[0], "{{ c1.full_name }}")
    set_para(kzs.paragraphs[2], "{{ c1.residential_address }}, {{ c1.phone }},")
    set_para(kzs.paragraphs[3], "{{ c1.email }}")
    set_run_text(kzs.paragraphs[6], 0, "ЖСН {{ c1.iin }}")
    set_para(kzs.paragraphs[7], "{{ c1.bank_name }}")
    set_para(kzs.paragraphs[8], "{{ c1.bank_account }}")
    set_para(kzs.paragraphs[10], "{% if second_insurer %}{{ c2.full_name }}{% endif %}")
    set_para(kzs.paragraphs[12], "{% if second_insurer %}{{ c2.residential_address }}, {{ c2.phone }},{% endif %}")
    set_para(kzs.paragraphs[13], "{% if second_insurer %}{{ c2.email }}{% endif %}")
    set_run_text(kzs.paragraphs[16], 0, "{% if second_insurer %}ЖСН {{ c2.iin }}{% endif %}")
    set_para(kzs.paragraphs[17], "{% if second_insurer %}{{ c2.bank_name }}{% endif %}")
    set_para(kzs.paragraphs[18], "{% if second_insurer %}{{ c2.bank_account }}{% endif %}")

    rus = t1.rows[4].cells[1]
    set_para(rus.paragraphs[1], "{{ c1.full_name }}")
    set_para(rus.paragraphs[3], "{{ c1.residential_address }}")
    set_para(rus.paragraphs[5], "{{ c1.phone }}, {{ c1.email }}")
    set_para(rus.paragraphs[7], "{{ c1.iin }}")
    set_para(rus.paragraphs[9], "{{ c1.bank_name }}")
    set_para(rus.paragraphs[11], "{{ c1.bank_account }}) (банка, № текущего счета)")
    set_para(rus.paragraphs[13], "{% if second_insurer %}{{ c2.full_name }}{% endif %}")
    set_para(rus.paragraphs[15], "{% if second_insurer %}{{ c2.residential_address }}{% endif %}")
    set_para(rus.paragraphs[17], "{% if second_insurer %}{{ c2.phone }}, {{ c2.email }}{% endif %}")
    set_para(rus.paragraphs[19], "{% if second_insurer %}{{ c2.iin }}{% endif %}")
    set_para(rus.paragraphs[21], "{% if second_insurer %}{{ c2.bank_name }}{% endif %}")
    set_para(rus.paragraphs[23], "{% if second_insurer %}{{ c2.bank_account }}) (банка, № текущего счета){% endif %}")

    # ======================================================================
    # TABLE 2 - payment schedule (appendix 1): turn row 1 into a docxtpl
    # table-row loop, drop the other 8 pre-printed blank rows.
    # ======================================================================
    # docxtpl row-loops need the {%tr for %} marker, the repeated data row
    # and the {%tr endfor %} marker on three *separate* table rows - a
    # marker row is fully consumed by the tag itself, any other content
    # placed in the same row is discarded. col0/col1 of this table are one
    # merged (gridSpan=2) cell, so python-docx returns the SAME cell object
    # twice per row - use distinct_cells() to only touch it once.
    row_for, row_data, row_endfor = t2.rows[1], t2.rows[2], t2.rows[3]

    dcells = distinct_cells(row_for)
    set_para(dcells[0].paragraphs[0], "{%tr for w in schedule %}")
    for c in dcells[1:]:
        set_para(c.paragraphs[0], "")

    dcells = distinct_cells(row_data)
    set_para(dcells[0].paragraphs[0], "{{ w.date_c1 }}")
    set_para(dcells[1].paragraphs[0], "{{ w.amount_c1 }}")
    set_para(dcells[2].paragraphs[0], "{{ w.buyout_c1 }}")
    set_para(dcells[3].paragraphs[0], "{{ w.date_c2 }}")
    set_para(dcells[4].paragraphs[0], "{{ w.amount_c2 }}")
    set_para(dcells[5].paragraphs[0], "{{ w.buyout_c2 }}")

    dcells = distinct_cells(row_endfor)
    set_para(dcells[0].paragraphs[0], "{%tr endfor %}")
    for c in dcells[1:]:
        set_para(c.paragraphs[0], "")

    # remaining pre-printed blank rows are no longer needed once the loop
    # can grow the table to any length - remove them to avoid stray rows
    tbl = t2._tbl
    for r in list(t2.rows[4:]):
        tbl.remove(r._tr)

    # ======================================================================
    # TABLE 3 - "Информация о застрахованном" structured summary appendix
    # ======================================================================
    mapping = {
        3: "{{ c1.full_name }}",
        4: "{{ c1.birth_date }}, {{ summary.age_c1 }} {{ summary.age_c1_label }}",
        5: "{{ summary.gender_c1 }}",
        7: "{% if second_insurer %}{{ c2.full_name }}{% endif %}",
        8: "{% if second_insurer %}{{ c2.birth_date }}, {{ summary.age_c2 }} {{ summary.age_c2_label }}{% endif %}",
        9: "{% if second_insurer %}{{ summary.gender_c2 }}{% endif %}",
        11: "{{ summary.contract_basis }}",
        12: "{{ premium.total }}",
        13: "{{ premium.premium_c1 }}",
        14: "{% if second_insurer %}{{ premium.premium_c2 }}{% endif %}",
        15: "{{ premium.other_org_c1 }}",
        16: "{{ premium.enpf_c1 }}",
        17: "{{ premium.own_c1 }}",
        18: "{% if second_insurer %}{{ premium.other_org_c2 }}{% endif %}",
        19: "{% if second_insurer %}{{ premium.enpf_c2 }}{% endif %}",
        20: "{% if second_insurer %}{{ premium.own_c2 }}{% endif %}",
        21: "{{ summary.term_c1 }}",
        22: "{% if second_insurer %}{{ summary.term_c2 }}{% endif %}",
        25: "{{ summary.effective_rate }}",
        26: "{{ summary.indexation_rate }}",
        27: "{{ summary.cost_ratio_premium }}",
        28: "{{ summary.cost_ratio_payment }}",
        29: "{{ summary.pv_factor }}",
        30: "{{ summary.pv_factor_costs }}",
        31: "{{ first_payment.c1 }}",
        32: "{% if second_insurer %}{{ first_payment.c2 }}{% endif %}",
    }
    for row_idx, tag in mapping.items():
        cell = t3.rows[row_idx].cells[1]
        set_para(cell.paragraphs[0], tag)

    # rows 23 & 24 - guaranteed-payment period ranges (4-paragraph blanks)
    row23 = t3.rows[23].cells[1]
    set_para(row23.paragraphs[0], 'с «{{ guarantee.c1_from_day }}»')
    set_para(row23.paragraphs[1], '{{ guarantee.c1_from_month }}')
    set_para(row23.paragraphs[2], '{{ guarantee.c1_from_year }} года по «{{ guarantee.c1_to_day }}» {{ guarantee.c1_to_month }}')
    set_para(row23.paragraphs[3], '{{ guarantee.c1_to_year }} года')
    row24 = t3.rows[24].cells[1]
    set_para(row24.paragraphs[0], '{% if second_insurer %}с «{{ guarantee.c2_from_day }}»{% endif %}')
    set_para(row24.paragraphs[1], '{% if second_insurer %}{{ guarantee.c2_from_month }}{% endif %}')
    set_para(row24.paragraphs[2], '{% if second_insurer %}{{ guarantee.c2_from_year }} года по «{{ guarantee.c2_to_day }}» {{ guarantee.c2_to_month }}{% endif %}')
    set_para(row24.paragraphs[3], '{% if second_insurer %}{{ guarantee.c2_to_year }} года{% endif %}')

    DST.parent.mkdir(parents=True, exist_ok=True)
    d.save(DST)
    print(f"Saved annotated contract template -> {DST}")


if __name__ == "__main__":
    main()
